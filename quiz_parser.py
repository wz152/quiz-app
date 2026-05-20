import re
import fitz
try:
    import pytesseract
    HAS_TESSERACT = True
except ImportError:
    pytesseract = None
    HAS_TESSERACT = False
from PIL import Image
from docx import Document
import io

def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Word文档解析失败: {str(e)}")

def extract_text_from_docx_bytes(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Word文档解析失败: {str(e)}")

def extract_text_from_pdf(filepath):
    try:
        doc = fitz.open(filepath)
        text = ""
        ocr_attempted = False
        
        for page_num in range(len(doc)):
            page = doc[page_num]

            text_blocks = page.get_text("blocks")
            has_text = any(b[4].strip() for b in text_blocks if len(b) >= 5 and isinstance(b[4], str))
            page_text = page.get_text("text")

            if has_text and page_text.strip():
                text += page_text + "\n"
            else:
                if HAS_TESSERACT:
                    try:
                        pix = page.get_pixmap(dpi=300)
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        ocr_text = pytesseract.image_to_string(img, lang='chi_sim+eng', config='--psm 6')
                        if ocr_text.strip():
                            text += ocr_text + "\n"
                            ocr_attempted = True
                    except Exception:
                        continue
        
        doc.close()

        if not text.strip():
            if ocr_attempted:
                raise Exception("PDF 文本和 OCR 均未能提取到内容，请检查文件是否可读")
            raise Exception("PDF 可能是扫描件且 Tesseract OCR 未安装，请安装 Tesseract 后再试")

        return text
    except Exception as e:
        if isinstance(e, Exception) and "PDF文档解析失败" not in str(e):
            raise Exception(f"PDF文档解析失败: {str(e)}")
        raise

def extract_text_from_pdf_bytes(file_bytes):
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        ocr_attempted = False
        
        for page_num in range(len(doc)):
            page = doc[page_num]

            text_blocks = page.get_text("blocks")
            has_text = any(b[4].strip() for b in text_blocks if len(b) >= 5 and isinstance(b[4], str))
            page_text = page.get_text("text")

            if has_text and page_text.strip():
                text += page_text + "\n"
            else:
                if HAS_TESSERACT:
                    try:
                        pix = page.get_pixmap(dpi=300)
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        ocr_text = pytesseract.image_to_string(img, lang='chi_sim+eng', config='--psm 6')
                        if ocr_text.strip():
                            text += ocr_text + "\n"
                            ocr_attempted = True
                    except Exception:
                        continue
        
        doc.close()

        if not text.strip():
            if ocr_attempted:
                raise Exception("PDF 文本和 OCR 均未能提取到内容，请检查文件是否可读")
            raise Exception("PDF 可能是扫描件且 Tesseract OCR 未安装，请安装 Tesseract 后再试")

        return text
    except Exception as e:
        if isinstance(e, Exception) and "PDF文档解析失败" not in str(e):
            raise Exception(f"PDF文档解析失败: {str(e)}")
        raise

def parse_questions(text):
    questions = []
    
    pattern = r'(\d+)\s*[.．、]\s*\(([^)]+)\)\s*'
    matches = list(re.finditer(pattern, text))
    
    for i, match in enumerate(matches):
        question_id = int(match.group(1))
        question_type = match.group(2).strip()
        
        start_pos = match.end()
        end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start_pos:end_pos].strip()
        
        question_data = {
            "id": question_id,
            "type": question_type,
            "question": "",
            "options": {},
            "correct_answer": "",
            "user_answer": ""
        }
        
        lines = body.split('\n')
        question_lines = []
        option_pattern = r'^([A-E])[.．、]\s*(.+)$'
        answer_pattern = r'正确答案[：:]\s*(.+)$'
        user_answer_pattern = r'我的答案[：:]\s*(.+)$'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            option_match = re.match(option_pattern, line)
            if option_match:
                option_key = option_match.group(1).upper()
                option_value = option_match.group(2).strip()
                question_data["options"][option_key] = option_value
            else:
                answer_match = re.search(answer_pattern, line)
                if answer_match:
                    question_data["correct_answer"] = answer_match.group(1).strip()
                    continue
                
                user_answer_match = re.search(user_answer_pattern, line)
                if user_answer_match:
                    question_data["user_answer"] = user_answer_match.group(1).strip()
                    continue
                
                question_lines.append(line)
        
        question_data["question"] = " ".join(question_lines).strip()
        
        if question_type == "填空题":
            bracket_pattern = r'【([^】]+)】'
            bracket_matches = re.findall(bracket_pattern, question_data["question"])
            if bracket_matches:
                question_data["correct_answer"] = bracket_matches[0]
                question_data["question"] = re.sub(bracket_pattern, "【  】", question_data["question"])
        
        questions.append(question_data)
    
    return {
        "questions": questions,
        "total": len(questions)
    }